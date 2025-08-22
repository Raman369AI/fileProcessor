"""
Universal File Processor
Reads multiple file formats and extracts content with table detection.
Minimal dependencies, optimized for resource efficiency.

ML Enhancement Opportunities:
- Document classification using pre-trained transformers
- Layout analysis with computer vision models (LayoutLM, DETR)
- Table detection using object detection models (YOLO, R-CNN)
- Text extraction quality improvement with post-processing models
- Semantic table understanding with TAPAS or similar models
- Content summarization using language models
- Entity extraction and relationship mapping
"""

import os
import re
import json
import csv
import io
import struct
import tempfile
import requests
from pathlib import Path
from typing import Dict, List, Tuple, Optional, Any
from dataclasses import dataclass
from datetime import datetime

try:
    import pandas as pd
    HAS_PANDAS = True
except ImportError:
    HAS_PANDAS = False

try:
    from PIL import Image
    import pytesseract
    HAS_OCR = True
except ImportError:
    HAS_OCR = False

try:
    import fitz  # PyMuPDF
    HAS_PYMUPDF = True
except ImportError:
    HAS_PYMUPDF = False

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import extract_msg
    HAS_EXTRACT_MSG = True
except ImportError:
    HAS_EXTRACT_MSG = False



@dataclass
class ExtractedContent:
    """
    Container for extracted file content
    
    ML Enhancement Opportunities:
    - Add confidence scores for extracted text and tables
    - Include bounding box coordinates for visual elements
    - Store feature embeddings for semantic search
    - Add classification labels (e.g., invoice, report, email)
    """
    text: str
    tables: List[List[List[str]]]
    metadata: Dict[str, Any]
    file_type: str


class AttachmentReader:
    """
    Modular attachment reader with custom processing functions for different file types
    
    ML Enhancement Opportunities:
    - Automatic attachment classification using content analysis
    - Smart extraction based on attachment importance scoring
    - Content-aware processing pipeline selection
    - Anomaly detection for suspicious attachments
    """
    
    def __init__(self, file_processor_instance=None):
        """
        Initialize attachment reader
        
        Args:
            file_processor_instance: Instance of FileProcessor for content processing
        """
        self.file_processor = file_processor_instance
        self.custom_processors = {}
        self.supported_types = {
            '.pdf': self._process_pdf_attachment,
            '.docx': self._process_docx_attachment,
            '.xlsx': self._process_excel_attachment,
            '.xls': self._process_excel_attachment,
            '.csv': self._process_csv_attachment,
            '.txt': self._process_text_attachment,
            '.jpg': self._process_image_attachment,
            '.jpeg': self._process_image_attachment,
            '.png': self._process_image_attachment,
            '.tiff': self._process_image_attachment,
            '.bmp': self._process_image_attachment
        }
    
    def register_custom_processor(self, file_extension: str, processor_function):
        """
        Register a custom processor function for specific file types
        
        Args:
            file_extension: File extension (e.g., '.pdf', '.docx')
            processor_function: Function that takes (file_path, metadata) and returns processed content
        """
        self.custom_processors[file_extension.lower()] = processor_function
    
    def read_attachment(self, attachment_data: bytes, filename: str, metadata: Dict[str, Any] = None) -> Dict[str, Any]:
        """
        Read and process attachment based on file type
        
        Args:
            attachment_data: Raw attachment bytes
            filename: Original filename with extension
            metadata: Additional metadata about the attachment
            
        Returns:
            Dict containing processed attachment content and metadata
            
        ML Enhancement Opportunities:
        - Use deep learning for file type detection beyond extensions
        - Implement content-based processing selection
        - Add confidence scoring for processing results
        """
        if metadata is None:
            metadata = {}
        
        result = {
            "filename": filename,
            "file_size": len(attachment_data),
            "processed_content": None,
            "processing_method": "none",
            "errors": [],
            "metadata": metadata
        }
        
        try:
            # Create temporary file for processing
            file_ext = os.path.splitext(filename)[1].lower()
            
            with tempfile.NamedTemporaryFile(suffix=file_ext, delete=False) as temp_file:
                temp_file.write(attachment_data)
                temp_file_path = temp_file.name
            
            try:
                # Check for custom processor first
                if file_ext in self.custom_processors:
                    result["processed_content"] = self.custom_processors[file_ext](
                        temp_file_path, metadata
                    )
                    result["processing_method"] = "custom"
                
                # Use built-in processors
                elif file_ext in self.supported_types:
                    result["processed_content"] = self.supported_types[file_ext](
                        temp_file_path, metadata
                    )
                    result["processing_method"] = "built-in"
                
                else:
                    # Try to process as text for unknown types
                    result["processed_content"] = self._process_unknown_attachment(
                        temp_file_path, metadata
                    )
                    result["processing_method"] = "fallback"
                    
            finally:
                # Clean up temporary file
                if os.path.exists(temp_file_path):
                    os.unlink(temp_file_path)
                    
        except Exception as e:
            result["errors"].append(f"Processing error: {str(e)}")
        
        return result
    
    def _process_pdf_attachment(self, file_path: str, metadata: Dict[str, Any]) -> ExtractedContent:
        """
        Default PDF processing - can be overridden with custom processor
        
        ML Enhancement Opportunities:
        - Document classification for PDF content routing
        - Form field detection and extraction
        - Invoice/receipt specific processing
        - Table structure understanding with specialized models
        """
        if self.file_processor:
            return self.file_processor._process_pdf(file_path)
        else:
            # Fallback processing without FileProcessor
            return self._basic_pdf_processing(file_path)
    
    def _basic_pdf_processing(self, file_path: str) -> ExtractedContent:
        """Basic PDF processing without full FileProcessor"""
        text = ""
        tables = []
        metadata = {"pages": 0}
        
        if HAS_PYMUPDF:
            try:
                doc = fitz.open(file_path)
                metadata["pages"] = len(doc)
                
                for page in doc:
                    text += page.get_text() + "\n"
                
                doc.close()
                
                # Basic table detection
                table_extractor = TableExtractor()
                tables = table_extractor.detect_table_patterns(text)
                
            except Exception as e:
                text = f"Error processing PDF: {str(e)}"
        else:
            text = "PyMuPDF not available for PDF processing"
        
        return ExtractedContent(
            text=text,
            tables=tables,
            metadata=metadata,
            file_type="pdf_attachment"
        )
    
    def _process_docx_attachment(self, file_path: str, metadata: Dict[str, Any]) -> ExtractedContent:
        """Process Word document attachments"""
        if self.file_processor:
            return self.file_processor._process_docx(file_path)
        else:
            return ExtractedContent(
                text="DOCX processing requires FileProcessor instance",
                tables=[],
                metadata=metadata,
                file_type="docx_attachment"
            )
    
    def _process_excel_attachment(self, file_path: str, metadata: Dict[str, Any]) -> ExtractedContent:
        """Process Excel attachments"""
        if self.file_processor:
            return self.file_processor._process_excel(file_path)
        else:
            return ExtractedContent(
                text="Excel processing requires FileProcessor instance",
                tables=[],
                metadata=metadata,
                file_type="excel_attachment"
            )
    
    def _process_csv_attachment(self, file_path: str, metadata: Dict[str, Any]) -> ExtractedContent:
        """Process CSV attachments"""
        if self.file_processor:
            return self.file_processor._process_csv(file_path)
        else:
            return ExtractedContent(
                text="CSV processing requires FileProcessor instance",
                tables=[],
                metadata=metadata,
                file_type="csv_attachment"
            )
    
    def _process_text_attachment(self, file_path: str, metadata: Dict[str, Any]) -> ExtractedContent:
        """Process text attachments"""
        if self.file_processor:
            return self.file_processor._process_text(file_path)
        else:
            try:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    text = f.read()
                return ExtractedContent(
                    text=text,
                    tables=[],
                    metadata=metadata,
                    file_type="text_attachment"
                )
            except Exception as e:
                return ExtractedContent(
                    text=f"Error reading text file: {str(e)}",
                    tables=[],
                    metadata=metadata,
                    file_type="text_attachment"
                )
    
    def _process_image_attachment(self, file_path: str, metadata: Dict[str, Any]) -> ExtractedContent:
        """Process image attachments with OCR"""
        if self.file_processor:
            return self.file_processor._process_image(file_path)
        else:
            return ExtractedContent(
                text="Image processing requires FileProcessor instance with OCR",
                tables=[],
                metadata=metadata,
                file_type="image_attachment"
            )
    
    def _process_unknown_attachment(self, file_path: str, metadata: Dict[str, Any]) -> ExtractedContent:
        """Fallback processing for unknown file types"""
        try:
            # Try to read as text
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()
            
            # Basic cleanup
            text = ''.join(char for char in text if char.isprintable() or char in '\n\r\t')
            
            return ExtractedContent(
                text=text,
                tables=[],
                metadata=metadata,
                file_type="unknown_attachment"
            )
        except Exception as e:
            return ExtractedContent(
                text=f"Unable to process unknown file type: {str(e)}",
                tables=[],
                metadata=metadata,
                file_type="unknown_attachment"
            )




class OutlookMsgParser:
    """
    Enhanced MSG file parser for Outlook emails
    
    ML Enhancement Opportunities:
    - Email thread detection using NLP models
    - Automatic categorization with BERT classification
    - Sentiment analysis for email content
    - Named entity recognition for contacts and organizations
    - Priority classification using machine learning
    """
    
    @staticmethod
    def parse_msg_file(file_path: str) -> Dict[str, Any]:
        """
        Parse MSG file and extract email components
        
        ML Enhancement Opportunities:
        - Use transformer models for email parsing and understanding
        - Implement attachment content analysis with multi-modal models
        - Add email signature detection using pattern recognition
        - Use NER for contact information extraction
        """
        msg_data = {
            "subject": "",
            "sender": "",
            "recipients": [],
            "date": "",
            "body": "",
            "attachments": [],
            "headers": {},
            "importance": "normal"
        }
        
        if HAS_EXTRACT_MSG:
            try:
                msg = extract_msg.Message(file_path)
                
                msg_data["subject"] = msg.subject or ""
                msg_data["sender"] = msg.sender or ""
                msg_data["date"] = str(msg.date) if msg.date else ""
                msg_data["body"] = msg.body or ""
                
                # Extract recipients
                if msg.to:
                    msg_data["recipients"].extend(msg.to.split(';'))
                if msg.cc:
                    msg_data["recipients"].extend(msg.cc.split(';'))
                if msg.bcc:
                    msg_data["recipients"].extend(msg.bcc.split(';'))
                
                # Clean recipients
                msg_data["recipients"] = [r.strip() for r in msg_data["recipients"] if r.strip()]
                
                # Extract attachments info
                for attachment in msg.attachments:
                    if hasattr(attachment, 'longFilename') and attachment.longFilename:
                        msg_data["attachments"].append({
                            "filename": attachment.longFilename,
                            "size": getattr(attachment, 'size', 0)
                        })
                
                # Extract headers
                if hasattr(msg, 'header'):
                    msg_data["headers"] = msg.header
                
                # Extract importance
                if hasattr(msg, 'importance'):
                    importance_map = {0: "low", 1: "normal", 2: "high"}
                    msg_data["importance"] = importance_map.get(msg.importance, "normal")
                
                msg.close()
                
            except Exception as e:
                # Fallback to basic parsing
                msg_data = OutlookMsgParser._basic_msg_parse(file_path)
                msg_data["parse_error"] = str(e)
        else:
            # Fallback to basic parsing
            msg_data = OutlookMsgParser._basic_msg_parse(file_path)
        
        return msg_data
    
    @staticmethod
    def _basic_msg_parse(file_path: str) -> Dict[str, Any]:
        """
        Basic MSG file parsing without external libraries
        
        ML Enhancement Opportunities:
        - Implement deep learning-based email structure detection
        - Use sequence models for email field extraction
        - Add fuzzy matching for partial email recovery
        """
        msg_data = {
            "subject": "",
            "sender": "",
            "recipients": [],
            "date": "",
            "body": "",
            "attachments": [],
            "headers": {},
            "importance": "normal",
            "extraction_method": "basic"
        }
        
        try:
            with open(file_path, 'rb') as file:
                content = file.read()
                
                # Try to extract readable text
                text_content = content.decode('utf-8', errors='ignore')
                
                # Basic pattern matching for email fields
                subject_match = re.search(r'Subject:\s*([^\r\n]+)', text_content, re.IGNORECASE)
                if subject_match:
                    msg_data["subject"] = subject_match.group(1).strip()
                
                from_match = re.search(r'From:\s*([^\r\n]+)', text_content, re.IGNORECASE)
                if from_match:
                    msg_data["sender"] = from_match.group(1).strip()
                
                to_match = re.search(r'To:\s*([^\r\n]+)', text_content, re.IGNORECASE)
                if to_match:
                    msg_data["recipients"] = [r.strip() for r in to_match.group(1).split(';')]
                
                date_match = re.search(r'Date:\s*([^\r\n]+)', text_content, re.IGNORECASE)
                if date_match:
                    msg_data["date"] = date_match.group(1).strip()
                
                # Extract body (everything after headers)
                body_start = text_content.find('\n\n')
                if body_start > 0:
                    potential_body = text_content[body_start:].strip()
                    # Clean up body text
                    msg_data["body"] = re.sub(r'[^\x20-\x7E\n\r\t]', '', potential_body)
                
                msg_data["file_size"] = len(content)
                
        except Exception as e:
            msg_data["parse_error"] = str(e)
            msg_data["body"] = f"Error parsing MSG file: {str(e)}"
        
        return msg_data
    
    @staticmethod
    def read_all_attachments(file_path: str, email_groups: List[str] = None, 
                           attachment_reader: 'AttachmentReader' = None, 
                           file_types: List[str] = None) -> Dict[str, Any]:
        """
        Read and process all attachments from MSG file using modular attachment reader
        
        Args:
            file_path: Path to MSG file
            email_groups: Filter by specific email groups/senders
            attachment_reader: AttachmentReader instance for processing
            file_types: List of file extensions to process (e.g., ['.pdf', '.docx'])
            
        Returns:
            Dict containing attachment info and processed content
            
        ML Enhancement Opportunities:
        - Classify attachment importance using content analysis
        - Detect document types using machine learning
        - Extract structured data using specialized models
        """
        extraction_result = {
            "attachments": [],
            "processed_content": [],
            "extraction_errors": [],
            "email_info": {},
            "summary": {}
        }
        
        try:
            # First parse the email to check sender/groups
            msg_data = OutlookMsgParser.parse_msg_file(file_path)
            extraction_result["email_info"] = {
                "subject": msg_data.get("subject", ""),
                "sender": msg_data.get("sender", ""),
                "date": msg_data.get("date", "")
            }
            
            # Check if email is from specified groups
            if email_groups:
                sender = msg_data.get("sender", "").lower()
                group_match = any(group.lower() in sender for group in email_groups)
                if not group_match:
                    extraction_result["extraction_errors"].append(
                        f"Email sender '{sender}' not in specified groups: {email_groups}"
                    )
                    return extraction_result
            
            if not HAS_EXTRACT_MSG:
                extraction_result["extraction_errors"].append(
                    "extract-msg library not available. Install with: pip install extract-msg"
                )
                return extraction_result
            
            # Initialize attachment reader if not provided
            if attachment_reader is None:
                attachment_reader = AttachmentReader()
            
            # Process MSG file for attachments
            msg = extract_msg.Message(file_path)
            
            processed_count = 0
            total_attachments = 0
            
            for attachment in msg.attachments:
                if hasattr(attachment, 'longFilename') and attachment.longFilename:
                    filename = attachment.longFilename
                    file_ext = os.path.splitext(filename)[1].lower()
                    total_attachments += 1
                    
                    # Check if we should process this file type
                    if file_types and file_ext not in file_types:
                        continue
                    
                    try:
                        # Get attachment data as bytes
                        attachment_data = attachment.data
                        
                        # Create attachment metadata
                        attachment_metadata = {
                            "filename": filename,
                            "size": getattr(attachment, 'size', len(attachment_data)),
                            "email_subject": msg_data.get("subject", ""),
                            "email_sender": msg_data.get("sender", ""),
                            "email_date": msg_data.get("date", "")
                        }
                        
                        # Process attachment using modular reader
                        processed_attachment = attachment_reader.read_attachment(
                            attachment_data, filename, attachment_metadata
                        )
                        
                        extraction_result["attachments"].append({
                            "filename": filename,
                            "file_type": file_ext,
                            "size": attachment_metadata["size"],
                            "processing_method": processed_attachment["processing_method"]
                        })
                        
                        extraction_result["processed_content"].append(processed_attachment)
                        processed_count += 1
                        
                    except Exception as e:
                        extraction_result["extraction_errors"].append(
                            f"Error processing attachment {filename}: {str(e)}"
                        )
            
            msg.close()
            
            # Build summary
            extraction_result["summary"] = {
                "total_attachments": total_attachments,
                "processed_attachments": processed_count,
                "file_types_processed": list(set(att["file_type"] for att in extraction_result["attachments"])),
                "processing_success_rate": f"{processed_count}/{total_attachments}" if total_attachments > 0 else "0/0"
            }
            
        except Exception as e:
            extraction_result["extraction_errors"].append(f"General extraction error: {str(e)}")
        
        return extraction_result


class TableExtractor:
    """
    Extract tables from text using pattern recognition
    
    ML Enhancement Opportunities:
    - Replace rule-based detection with deep learning models
    - Use computer vision for table boundary detection
    - Implement BERT-based cell classification
    - Add table structure understanding (header detection, merged cells)
    - Use graph neural networks for table relationship modeling
    """
    
    @staticmethod
    def detect_table_patterns(text: str) -> List[List[List[str]]]:
        """
        Detect table-like structures in text using heuristic patterns
        
        Current approach: Rule-based pattern matching
        
        ML Enhancement Opportunities:
        - Replace with transformer-based sequence labeling (BERT, RoBERTa)
        - Use named entity recognition for column header detection
        - Implement attention mechanisms to understand cell relationships
        - Add fuzzy matching for inconsistent formatting
        - Use clustering algorithms to group similar table structures
        """
        tables = []
        lines = text.strip().split('\n')
        
        # Pattern 1: Tab or multiple space separated values
        potential_table = []
        for line in lines:
            # Check if line has multiple columns (tab or 2+ spaces)
            if '\t' in line or re.search(r'\s{2,}', line):
                cols = re.split(r'\t|\s{2,}', line.strip())
                if len(cols) > 1:
                    potential_table.append(cols)
                elif potential_table:
                    # End of table
                    if len(potential_table) > 1:
                        tables.append(potential_table)
                    potential_table = []
            elif potential_table:
                # End of table
                if len(potential_table) > 1:
                    tables.append(potential_table)
                potential_table = []
        
        # Add last table if exists
        if potential_table and len(potential_table) > 1:
            tables.append(potential_table)
        
        # Pattern 2: Pipe separated values
        pipe_table = []
        for line in lines:
            if '|' in line and line.count('|') > 1:
                cols = [col.strip() for col in line.split('|') if col.strip()]
                if cols:
                    pipe_table.append(cols)
            elif pipe_table:
                if len(pipe_table) > 1:
                    tables.append(pipe_table)
                pipe_table = []
        
        if pipe_table and len(pipe_table) > 1:
            tables.append(pipe_table)
        
        return tables
    
    @staticmethod
    def extract_from_coordinates(text: str) -> List[List[List[str]]]:
        """
        Extract tables using coordinate-based detection
        
        Current approach: Simple position-based grouping
        
        ML Enhancement Opportunities:
        - Use computer vision models to detect table regions in document images
        - Implement YOLO or R-CNN for table detection
        - Use OCR with layout analysis (Azure Form Recognizer, AWS Textract)
        - Apply clustering algorithms (DBSCAN, K-means) for cell grouping
        - Use graph neural networks to model spatial relationships
        - Implement reinforcement learning for optimal table boundary detection
        """
        lines = text.split('\n')
        tables = []
        
        # Look for lines with consistent column positions
        column_positions = {}
        for line_idx, line in enumerate(lines):
            words = []
            for match in re.finditer(r'\S+', line):
                words.append((match.start(), match.group()))
            
            if len(words) > 1:
                positions = tuple(word[0] for word in words)
                if positions not in column_positions:
                    column_positions[positions] = []
                column_positions[positions].append((line_idx, [word[1] for word in words]))
        
        # Extract tables from consistent column positions
        for positions, lines_data in column_positions.items():
            if len(lines_data) > 1:  # At least 2 rows for a table
                table = [line_data[1] for line_data in lines_data]
                tables.append(table)
        
        return tables


class FileProcessor:
    """
    Main file processing class with modular readers
    
    ML Enhancement Opportunities:
    - Add document classification pipeline using BERT/DistilBERT
    - Implement multi-modal processing (text + images) with CLIP
    - Use few-shot learning for new file format adaptation
    - Add active learning for improving extraction accuracy
    - Implement federated learning for privacy-preserving model updates
    """
    
    def __init__(self):
        self.table_extractor = TableExtractor()
        self.outlook_parser = OutlookMsgParser()
        
        # Initialize AttachmentReader with this FileProcessor instance
        self.attachment_reader = AttachmentReader(file_processor_instance=self)
    
    def process_file(self, file_path: str) -> ExtractedContent:
        """
        Process any supported file type
        
        ML Enhancement Opportunities:
        - Add automatic file type detection using content analysis
        - Implement confidence scoring for processing results
        - Use ensemble methods combining multiple extraction approaches
        - Add error correction using language models
        """
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_ext = path.suffix.lower()
        
        # Route to appropriate processor
        if file_ext == '.pdf':
            return self._process_pdf(file_path)
        elif file_ext in ['.jpg', '.jpeg', '.png', '.tiff', '.bmp']:
            return self._process_image(file_path)
        elif file_ext == '.docx':
            return self._process_docx(file_path)
        elif file_ext in ['.xlsx', '.xls']:
            return self._process_excel(file_path)
        elif file_ext == '.csv':
            return self._process_csv(file_path)
        elif file_ext == '.msg':
            return self._process_outlook(file_path)
        elif file_ext == '.txt':
            return self._process_text(file_path)
        else:
            # Try as text file
            return self._process_text(file_path)
    
    def _process_pdf(self, file_path: str) -> ExtractedContent:
        """
        Process PDF files using PyMuPDF with OCR fallback
        
        ML Enhancement Opportunities:
        - Use LayoutLM for document understanding and structure recognition
        - Implement PDF layout analysis with computer vision models
        - Add reading order detection using sequence models
        - Use BERT for content classification and entity extraction
        - Implement form field detection using object detection models
        - Add table detection using specialized models like TableNet
        """
        text = ""
        tables = []
        metadata = {"pages": 0, "images": 0, "tables_detected": 0}
        
        if HAS_PYMUPDF:
            try:
                doc = fitz.open(file_path)
                metadata["pages"] = len(doc)
                
                for page_num in range(len(doc)):
                    page = doc[page_num]
                    
                    # Extract text
                    page_text = page.get_text()
                    if page_text.strip():
                        text += f"--- Page {page_num + 1} ---\n"
                        text += page_text + "\n"
                    
                    # Extract tables using PyMuPDF's table detection
                    try:
                        page_tables = page.find_tables()
                        metadata["tables_detected"] += len(page_tables)
                        
                        for table in page_tables:
                            table_data = table.extract()
                            if table_data:
                                # Clean table data
                                cleaned_table = []
                                for row in table_data:
                                    cleaned_row = [str(cell).strip() if cell else "" for cell in row]
                                    cleaned_table.append(cleaned_row)
                                if cleaned_table:
                                    tables.append(cleaned_table)
                    except Exception as table_error:
                        # Fallback to text-based table detection
                        pass
                    
                    # Count images
                    image_list = page.get_images()
                    metadata["images"] += len(image_list)
                    
                    # If no text extracted and OCR is available, try OCR
                    if not page_text.strip() and HAS_OCR:
                        try:
                            # Convert page to image for OCR
                            pix = page.get_pixmap()
                            img_data = pix.tobytes("png")
                            image = Image.open(io.BytesIO(img_data))
                            ocr_text = pytesseract.image_to_string(image)
                            if ocr_text.strip():
                                text += f"--- Page {page_num + 1} (OCR) ---\n"
                                text += ocr_text + "\n"
                        except Exception as ocr_error:
                            text += f"--- Page {page_num + 1} (OCR Failed) ---\n"
                
                doc.close()
                
                # If no tables found with PyMuPDF, fallback to pattern detection
                if not tables:
                    tables = self.table_extractor.detect_table_patterns(text)
                
            except Exception as e:
                text = f"Error processing PDF with PyMuPDF: {str(e)}"
                if HAS_OCR:
                    text += "\nAttempting OCR fallback..."
                    text += self._ocr_pdf_fallback(file_path)
        
        elif HAS_OCR:
            text = self._ocr_pdf_fallback(file_path)
        else:
            text = "PDF processing not available - install PyMuPDF (fitz) and/or pytesseract"
        
        return ExtractedContent(
            text=text,
            tables=tables,
            metadata=metadata,
            file_type="pdf"
        )
    
    def _ocr_pdf_fallback(self, file_path: str) -> str:
        """
        OCR processing for PDFs when PyMuPDF text extraction fails
        
        ML Enhancement Opportunities:
        - Use PaddleOCR or EasyOCR for better multilingual support
        - Implement text detection with EAST or CRAFT models
        - Add text recognition confidence scoring
        - Use post-processing with language models for error correction
        - Implement layout-aware OCR with PICK or LayoutLM
        - Add handwriting recognition capabilities with specialized models
        """
        if not HAS_OCR:
            return "OCR not available"
        
        try:
            # Try with pdf2image if available
            try:
                from pdf2image import convert_from_path
                pages = convert_from_path(file_path)
                text = ""
                for i, page in enumerate(pages):
                    text += f"--- Page {i + 1} (OCR) ---\n"
                    text += pytesseract.image_to_string(page) + "\n"
                return text
            except ImportError:
                # Alternative: Use PyMuPDF to convert to images
                if HAS_PYMUPDF:
                    doc = fitz.open(file_path)
                    text = ""
                    for page_num in range(len(doc)):
                        page = doc[page_num]
                        pix = page.get_pixmap()
                        img_data = pix.tobytes("png")
                        image = Image.open(io.BytesIO(img_data))
                        page_text = pytesseract.image_to_string(image)
                        text += f"--- Page {page_num + 1} (OCR) ---\n"
                        text += page_text + "\n"
                    doc.close()
                    return text
                else:
                    return "PDF to image conversion not available"
                    
        except Exception as e:
            return f"OCR error: {str(e)}"
    
    def _process_image(self, file_path: str) -> ExtractedContent:
        """
        Process image files with OCR
        
        ML Enhancement Opportunities:
        - Use YOLO or R-CNN for object detection and layout analysis
        - Implement scene text detection with specialized models (EAST, CRAFT)
        - Add image preprocessing with super-resolution models
        - Use vision transformers for document understanding
        - Implement table detection in images using TableNet or similar
        - Add handwriting recognition with specialized deep learning models
        """
        text = ""
        tables = []
        metadata = {}
        
        if HAS_OCR:
            try:
                image = Image.open(file_path)
                metadata = {
                    "width": image.width,
                    "height": image.height,
                    "mode": image.mode
                }
                text = pytesseract.image_to_string(image)
                tables = self.table_extractor.detect_table_patterns(text)
            except Exception as e:
                text = f"Error processing image: {str(e)}"
        else:
            text = "OCR not available - install pytesseract and PIL"
        
        return ExtractedContent(
            text=text,
            tables=tables,
            metadata=metadata,
            file_type="image"
        )
    
    def _process_docx(self, file_path: str) -> ExtractedContent:
        """
        Process Word documents
        
        ML Enhancement Opportunities:
        - Use transformer models for content summarization
        - Implement style and formatting analysis with computer vision
        - Add entity extraction using spaCy or Transformers
        - Use BERT for document classification and topic modeling
        - Implement table understanding with TAPAS or similar models
        - Add relationship extraction between document elements
        """
        text = ""
        tables = []
        metadata = {}
        
        if HAS_DOCX:
            try:
                doc = Document(file_path)
                
                # Extract paragraphs
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\n"
                
                # Extract tables
                for table in doc.tables:
                    table_data = []
                    for row in table.rows:
                        row_data = [cell.text.strip() for cell in row.cells]
                        table_data.append(row_data)
                    if table_data:
                        tables.append(table_data)
                
                metadata = {
                    "paragraphs": len(doc.paragraphs),
                    "tables": len(doc.tables)
                }
                
            except Exception as e:
                text = f"Error processing DOCX: {str(e)}"
        else:
            text = "DOCX processing not available - install python-docx"
        
        return ExtractedContent(
            text=text,
            tables=tables,
            metadata=metadata,
            file_type="docx"
        )
    
    def _process_excel(self, file_path: str) -> ExtractedContent:
        """
        Process Excel files
        
        ML Enhancement Opportunities:
        - Use machine learning for formula understanding and validation
        - Implement anomaly detection for data quality assessment
        - Add time series analysis for temporal data patterns
        - Use clustering for data categorization and grouping
        - Implement regression models for missing value imputation
        - Add natural language generation for data summaries
        """
        text = ""
        tables = []
        metadata = {}
        
        if HAS_PANDAS:
            try:
                # Read all sheets
                excel_file = pd.ExcelFile(file_path)
                metadata["sheets"] = excel_file.sheet_names
                
                for sheet_name in excel_file.sheet_names:
                    df = pd.read_excel(file_path, sheet_name=sheet_name)
                    
                    # Convert to table format
                    table_data = [df.columns.tolist()]
                    table_data.extend(df.values.tolist())
                    tables.append(table_data)
                    
                    # Add to text
                    text += f"Sheet: {sheet_name}\n"
                    text += df.to_string() + "\n\n"
                
            except Exception as e:
                text = f"Error processing Excel: {str(e)}"
        else:
            text = "Excel processing not available - install pandas"
        
        return ExtractedContent(
            text=text,
            tables=tables,
            metadata=metadata,
            file_type="excel"
        )
    
    def _process_csv(self, file_path: str) -> ExtractedContent:
        """
        Process CSV files
        
        ML Enhancement Opportunities:
        - Use machine learning for automatic delimiter detection
        - Implement data type inference using statistical models
        - Add outlier detection and anomaly analysis
        - Use clustering for data profiling and pattern discovery
        - Implement automated data quality assessment
        - Add predictive modeling for missing values
        """
        text = ""
        tables = []
        metadata = {}
        
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                # Detect delimiter
                sample = file.read(1024)
                file.seek(0)
                sniffer = csv.Sniffer()
                delimiter = sniffer.sniff(sample).delimiter
                
                reader = csv.reader(file, delimiter=delimiter)
                table_data = list(reader)
                
                if table_data:
                    tables.append(table_data)
                    
                    # Convert to text
                    for row in table_data:
                        text += delimiter.join(str(cell) for cell in row) + "\n"
                
                metadata = {
                    "rows": len(table_data),
                    "columns": len(table_data[0]) if table_data else 0,
                    "delimiter": delimiter
                }
                
        except Exception as e:
            text = f"Error processing CSV: {str(e)}"
        
        return ExtractedContent(
            text=text,
            tables=tables,
            metadata=metadata,
            file_type="csv"
        )
    
    def _process_outlook(self, file_path: str) -> ExtractedContent:
        """
        Process Outlook MSG files with enhanced parsing
        
        ML Enhancement Opportunities:
        - Use transformer models for email classification and sentiment analysis
        - Implement named entity recognition for contact and organization extraction
        - Add email thread understanding using sequence models
        - Use BERT for priority classification and urgency detection
        - Implement attachment analysis with multi-modal models
        - Add spam detection using ensemble methods
        """
        text = ""
        tables = []
        metadata = {}
        
        try:
            # Parse MSG file using enhanced parser
            msg_data = self.outlook_parser.parse_msg_file(file_path)
            
            # Build formatted text output
            text += f"Subject: {msg_data.get('subject', '')}\n"
            text += f"From: {msg_data.get('sender', '')}\n"
            
            if msg_data.get('recipients'):
                text += f"To: {'; '.join(msg_data['recipients'])}\n"
            
            text += f"Date: {msg_data.get('date', '')}\n"
            text += f"Importance: {msg_data.get('importance', 'normal')}\n"
            text += "\n" + "-" * 50 + "\n\n"
            text += msg_data.get('body', '')
            
            # Extract tables from email body
            if msg_data.get('body'):
                tables = self.table_extractor.detect_table_patterns(msg_data['body'])
            
            # Build metadata
            metadata = {
                "subject": msg_data.get('subject', ''),
                "sender": msg_data.get('sender', ''),
                "recipient_count": len(msg_data.get('recipients', [])),
                "date": msg_data.get('date', ''),
                "attachment_count": len(msg_data.get('attachments', [])),
                "importance": msg_data.get('importance', 'normal'),
                "file_size": msg_data.get('file_size', 0),
                "extraction_method": msg_data.get('extraction_method', 'enhanced'),
                "has_extract_msg": HAS_EXTRACT_MSG
            }
            
            # Add attachment info to metadata
            if msg_data.get('attachments'):
                metadata["attachments"] = msg_data['attachments']
            
            # Add any parsing errors to metadata
            if msg_data.get('parse_error'):
                metadata["parse_error"] = msg_data['parse_error']
                
        except Exception as e:
            text = f"Error processing MSG file: {str(e)}"
            metadata = {"error": str(e), "has_extract_msg": HAS_EXTRACT_MSG}
        
        return ExtractedContent(
            text=text,
            tables=tables,
            metadata=metadata,
            file_type="outlook"
        )
    
    def _process_text(self, file_path: str) -> ExtractedContent:
        """
        Process plain text files
        
        ML Enhancement Opportunities:
        - Use language detection models for multilingual support
        - Implement topic modeling with LDA or BERT-based approaches
        - Add sentiment analysis and emotion detection
        - Use named entity recognition for information extraction
        - Implement text summarization with transformer models
        - Add readability scoring and complexity analysis
        """
        text = ""
        tables = []
        metadata = {}
        
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                text = file.read()
                
            tables = self.table_extractor.detect_table_patterns(text)
            metadata = {
                "lines": text.count('\n'),
                "characters": len(text)
            }
            
        except Exception as e:
            text = f"Error processing text file: {str(e)}"
        
        return ExtractedContent(
            text=text,
            tables=tables,
            metadata=metadata,
            file_type="text"
        )
    
    def to_markdown(self, content: ExtractedContent) -> str:
        """
        Convert extracted content to markdown format
        
        ML Enhancement Opportunities:
        - Use natural language generation for intelligent formatting
        - Implement template selection based on content type classification
        - Add automatic heading generation using summarization models
        - Use language models for improved table caption generation
        """
        md = f"# File Content ({content.file_type.upper()})\n\n"
        
        # Add metadata
        if content.metadata:
            md += "## Metadata\n\n"
            for key, value in content.metadata.items():
                md += f"- **{key.title()}**: {value}\n"
            md += "\n"
        
        # Add text content
        if content.text.strip():
            md += "## Text Content\n\n"
            md += content.text + "\n\n"
        
        # Add tables
        if content.tables:
            md += "## Extracted Tables\n\n"
            for i, table in enumerate(content.tables, 1):
                md += f"### Table {i}\n\n"
                if table:
                    # Create markdown table
                    if len(table) > 0:
                        # Header
                        md += "| " + " | ".join(str(cell) for cell in table[0]) + " |\n"
                        md += "| " + " | ".join("---" for _ in table[0]) + " |\n"
                        
                        # Rows
                        for row in table[1:]:
                            md += "| " + " | ".join(str(cell) for cell in row) + " |\n"
                    md += "\n"
        
        return md
    
    def to_json(self, content: ExtractedContent) -> str:
        """
        Convert extracted content to JSON format
        
        ML Enhancement Opportunities:
        - Add schema validation using machine learning models
        - Implement automatic data type inference and conversion
        - Use natural language processing for field name standardization
        - Add confidence scores and uncertainty quantification
        """
        data = {
            "file_type": content.file_type,
            "metadata": content.metadata,
            "text": content.text,
            "tables": content.tables,
            "table_count": len(content.tables)
        }
        
        return json.dumps(data, indent=2, ensure_ascii=False)
    
    def read_attachments_from_msg(self, file_path: str, email_groups: List[str] = None, 
                                 file_types: List[str] = None, custom_processors: Dict[str, callable] = None) -> Dict[str, Any]:
        """
        Read and process attachments from MSG files using modular attachment reader
        
        Args:
            file_path: Path to MSG file
            email_groups: Filter by specific email groups/senders
            file_types: List of file extensions to process (e.g., ['.pdf', '.docx'])
            custom_processors: Dict of custom processing functions {'.ext': function}
            
        Returns:
            Dict containing extraction results and processed content
        """
        # Register custom processors if provided
        if custom_processors:
            for ext, processor in custom_processors.items():
                self.attachment_reader.register_custom_processor(ext, processor)
        
        return self.outlook_parser.read_all_attachments(
            file_path, email_groups, self.attachment_reader, file_types
        )
    
    def register_custom_attachment_processor(self, file_extension: str, processor_function):
        """
        Register a custom processor for specific attachment types
        
        Args:
            file_extension: File extension (e.g., '.pdf', '.docx')
            processor_function: Function that takes (file_path, metadata) and returns ExtractedContent
        """
        self.attachment_reader.register_custom_processor(file_extension, processor_function)
    
    def extract_pdf_attachments_from_msg(self, file_path: str, output_dir: str = None, email_groups: List[str] = None) -> Dict[str, Any]:
        """
        Legacy method - Extract and process PDF attachments from MSG files
        
        Args:
            file_path: Path to MSG file
            output_dir: Directory to save extracted PDFs (deprecated - attachments are processed in memory)
            email_groups: Filter by specific email groups/senders
            
        Returns:
            Dict containing extraction results and processed content
        """
        return self.read_attachments_from_msg(file_path, email_groups, ['.pdf'])
    


def main():
    """
    Example usage
    
    ML Enhancement Opportunities:
    - Add interactive CLI with intelligent file suggestions
    - Implement batch processing with progress prediction
    - Use reinforcement learning for user preference adaptation
    - Add real-time performance monitoring and optimization
    """
    processor = FileProcessor()
    
    # Example file processing
    file_path = input("Enter file path: ")
    
    try:
        content = processor.process_file(file_path)
        
        print("\n" + "="*50)
        print("MARKDOWN OUTPUT")
        print("="*50)
        print(processor.to_markdown(content))
        
        print("\n" + "="*50)
        print("JSON OUTPUT")
        print("="*50)
        print(processor.to_json(content))
        
    except Exception as e:
        print(f"Error: {e}")


if __name__ == "__main__":
    main()